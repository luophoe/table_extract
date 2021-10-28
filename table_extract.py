# table_extract.py
#
# Created on: 04/25/2021
#     Author: Anyka
#      		  Phoebe Luo
import re
from docx import Document

# open file and extract paragraphs and tables
doc = Document()
document = Document(r"C:\Users\anyka\Desktop\Snowbird3_GPIO模块设计文档_V1.0.5.docx")
paragraphs = document.paragraphs
tables = document.tables


# class to hold register information
class Register:
    name = None
    address = None
    dict_bit = {}
    dict_RW = {}
    dict_resetVal = {}
    dict_name = {}
    dict_desc = {}

    def __init__(self, name, address):
        self.name = name
        self.address = address

        self.dict_bit = {}
        self.dict_RW = {}
        self.dict_resetVal = {}
        self.dict_name = {}
        self.dict_desc = {}

    def getBitInfo(self, row, bit, RW, resetVal, name, desc):
        self.dict_bit[row] = bit
        self.dict_RW[row] = RW
        self.dict_resetVal[row] = resetVal
        self.dict_name[row] = name
        self.dict_desc[row] = desc


# function to do number system conversion
def numSysConv(text):
    num_sys = re.search(r'\’\s*(h|d|b)', text, re.M | re.I).group(1)
    if num_sys == "d":
        num_dec = re.search(r'\’\s*d(\d+)', text, re.M | re.I).group(1)
        num_bin = bin(int(num_dec))
        num_bin = num_bin.replace("0b", "b")
        text_bin = "'" + str(num_bin)
        return text_bin
    return text


# ------------------ Part 1. Extract Document Info ------------------
# match tables with keyword "Register" in the paragraph above
reg_dict = {}  # hold reg objects
reg_count = 0  # keep track of current reg count

for aPara in paragraphs:
    break_check = 0  # turn to 1 when a stored register is found
    searchText = re.search(r'Register', aPara.text, re.M | re.I)  # search keyword
    if searchText:
        specText = aPara.text.encode('utf-8').decode('utf-8')  # line with keyword
        # 1. match register name and address
        reg_addr = re.search(r'\(\s*(\w+)\s*\)', specText, re.M | re.I).group(1)
        reg_name = re.sub(r'\s*Register\s*\(\s*(\w+)\s*\)\s*', "", specText)
        # check if the register is already stored
        for i in range(reg_count):
            if reg_dict[i].address == reg_addr:
                break_check = 1
        if break_check == 0:
            obj_name = Register(reg_name, reg_addr)
            reg_dict[reg_count] = obj_name
            reg_count = reg_count + 1  # update reg_count
            # 2. find register bit information from corresponding table
            ele = aPara._p.getnext()
            row_count = 0  # keep track of the current row of the bit information is on
            while ele.tag != '' and ele.tag[-3:] != 'tbl':
                ele = ele.getnext()
            if ele.tag != '':
                for aTable in tables:
                    if aTable._tbl == ele:
                        for i in range(1, len(aTable.rows)):
                            bit = aTable.cell(i, 0).text
                            RW = aTable.cell(i, 1).text
                            resetVal = aTable.cell(i, 2).text
                            name = aTable.cell(i, 3).text
                            desc = aTable.cell(i, 4).text
                            obj_name.getBitInfo(row_count, bit, RW, resetVal, name, desc)
                            row_count = row_count + 1
                        break


# ------------------ Part 2. Writing Register Info ------------------
file = open(r"C:\Users\anyka\Desktop\ao_mem.ralf", "w+")
filename = file.name
filename = re.search(r'\\(\w*?)\.ralf', filename, re.M | re.I).group(1)

for i in range(len(reg_dict)):
    # 1. register ao_mem_operation {
    file.write("register " + reg_dict[i].name + " {" + "\n")
    for j in range(len(reg_dict[i].dict_RW) - 1, -1, -1):
        file.write("\n")
        # 2. field  operation_start {
        file.write("\t" + "field" + "\t" + reg_dict[i].dict_name[j] + " {" + "\n")
        # 3. bits   1;
        bits = re.search(r'(\d+)\s*\’', reg_dict[i].dict_resetVal[j], re.M | re.I).group(1)
        file.write("\t\t" + "bits" + "\t" + bits + ";" + "\n")
        # 4. access rw;
        if reg_dict[i].dict_RW[j] == "RW":
            access = "rw"
        elif reg_dict[i].dict_RW[j] == "RO":
            access = "ro"
        elif reg_dict[i].dict_RW[j] == "WO":
            access = "wo"
        else:
            access = ""
        file.write("\t\t" + "access" + "\t" + access + ";" + "\n")
        # 5. reset  'h0;
        reset_orig = re.search(r'\’\s*(.+)', reg_dict[i].dict_resetVal[j], re.M | re.I).group()
        reset = numSysConv(reset_orig)
        if re.match("’", reset):
            reset = reset.replace("’", "'")
        file.write("\t\t" + "reset" + "\t" + reset + ";" + "\n")
        file.write("\t" + "}\n")
        file.write("\n")
    file.write("\n}\n")
    file.write("\n")

file.write("block " + filename + " {\n")
bytes = str(len(reg_dict))
file.write("\t" + "bytes" + "\t" + bytes + ";\n")
file.write("\n")
for i in range(len(reg_dict)):
    file.write("\t" + "register" + "\t" + reg_dict[i].name + "\t" + "(" + reg_dict[i].name + "_bd" + ")" + "\t" + "@'h" + reg_dict[i].address[-2:] + ";\n")
file.write("\n}\n")